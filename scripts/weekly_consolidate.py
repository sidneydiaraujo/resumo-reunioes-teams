# -*- coding: utf-8 -*-
"""
weekly_consolidate.py
Lê os resumos .docx da semana atual no OneDrive e gera um relatório HTML
com todos os combinados e pendências consolidados.

Uso:
  python weekly_consolidate.py --output relatorio_semanal.html
"""

import sys
import json
import argparse
import tempfile
from pathlib import Path
from datetime import datetime, timedelta

TOKEN_CACHE_FILE = Path.home() / ".claude" / "onedrive_auth_cache.bin"
CLIENT_ID = "14d82eec-204b-4c2f-b7e8-296a70dab67e"
AUTHORITY = "https://login.microsoftonline.com/common"
SCOPES = ["Files.ReadWrite", "User.Read"]
ROOT_FOLDER = "Reuniões dos Times"


def ensure_deps():
    try:
        import msal, requests
        from docx import Document
    except ImportError:
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "msal", "requests", "python-docx", "-q"])


def get_token() -> str:
    import msal

    cache = msal.SerializableTokenCache()
    if TOKEN_CACHE_FILE.exists():
        cache.deserialize(TOKEN_CACHE_FILE.read_text(encoding="utf-8"))

    app = msal.PublicClientApplication(CLIENT_ID, authority=AUTHORITY, token_cache=cache)
    accounts = app.get_accounts()
    result = app.acquire_token_silent(SCOPES, account=accounts[0]) if accounts else None

    if not result or "access_token" not in result:
        raise RuntimeError("Token expirado. Execute onedrive_upload.py para re-autenticar.")

    TOKEN_CACHE_FILE.parent.mkdir(parents=True, exist_ok=True)
    TOKEN_CACHE_FILE.write_text(cache.serialize(), encoding="utf-8")
    return result["access_token"]


def get_week_range():
    today = datetime.now()
    monday = today - timedelta(days=today.weekday())
    sunday = monday + timedelta(days=6)
    return monday.replace(hour=0, minute=0, second=0), sunday.replace(hour=23, minute=59, second=59)


def list_week_files(token: str) -> list:
    import requests

    headers = {"Authorization": f"Bearer {token}", "Accept": "application/json"}
    week_start, week_end = get_week_range()
    files = []

    # Lista subpastas de Reuniões dos Times
    resp = requests.get(
        f"https://graph.microsoft.com/v1.0/me/drive/root:/{ROOT_FOLDER}:/children",
        headers=headers
    )
    if resp.status_code != 200:
        print(f"Aviso: não encontrou {ROOT_FOLDER}: {resp.status_code}")
        return files

    team_folders = resp.json().get("value", [])
    for folder in team_folders:
        if "folder" not in folder:
            continue
        folder_id = folder["id"]
        team_name = folder["name"]

        files_resp = requests.get(
            f"https://graph.microsoft.com/v1.0/me/drive/items/{folder_id}/children",
            headers=headers
        )
        for item in files_resp.json().get("value", []):
            if not item["name"].endswith(".docx"):
                continue
            modified = datetime.fromisoformat(item["lastModifiedDateTime"].replace("Z", "+00:00")).replace(tzinfo=None)
            if week_start <= modified <= week_end:
                files.append({
                    "name": item["name"],
                    "team": team_name,
                    "id": item["id"],
                    "modified": modified.strftime("%d/%m/%Y %H:%M")
                })

    return files


def download_and_parse_smart(doc) -> dict:
    """Parse documentos no novo formato inteligente (create_docx_smart.py)."""
    result = {
        "titulo": "", "data": "", "time": "", "participantes": [],
        "tipo_reuniao": "", "contexto": "",
        "secoes": [], "acoes_pendentes": [],
        "formato": "smart"
    }

    current_section = None
    current_items = []
    pending_bold = None

    def flush_section():
        if current_section and current_items:
            result["secoes"].append({"titulo": current_section, "itens": list(current_items)})

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or "═" in text or "─" in text or "RESUMO DE REUNIÃO" in text:
            continue

        # Metadados
        for label, key in [("Título:", "titulo"), ("Data:", "data"), ("Time:", "time"), ("Tipo:", "tipo_reuniao")]:
            if text.startswith(label):
                result[key] = text[len(label):].strip()
                break
        else:
            if text.startswith("Participantes:"):
                result["participantes"] = [p.strip() for p in text[len("Participantes:"):].split(",")]
            elif text.startswith("Resumo gerado em:"):
                flush_section()
                current_section = None
                current_items = []
            elif text == "AÇÕES PENDENTES":
                flush_section()
                current_section = "__acoes__"
                current_items = []
            elif text.isupper() and len(text) > 3:
                # Novo cabeçalho de seção
                flush_section()
                current_section = text.title()
                current_items = []
                pending_bold = None
            elif current_section == "__acoes__":
                # Itens de ações (podem ter linha de responsável/prazo abaixo)
                if text.startswith("•"):
                    item_text = text.lstrip("•").strip()
                    pending_bold = item_text
                elif pending_bold and ("Responsável:" in text or "Prazo:" in text):
                    resp = ""
                    prazo = ""
                    for part in text.split("  "):
                        part = part.strip()
                        if part.startswith("Responsável:"):
                            resp = part[len("Responsável:"):].strip()
                        elif part.startswith("Prazo:"):
                            prazo = part[len("Prazo:"):].strip()
                    result["acoes_pendentes"].append({"item": pending_bold, "responsavel": resp, "prazo": prazo})
                    pending_bold = None
                elif pending_bold:
                    result["acoes_pendentes"].append({"item": pending_bold, "responsavel": "", "prazo": ""})
                    pending_bold = None
            elif current_section:
                if text.startswith("•"):
                    item_text = text.lstrip("•").strip()
                    pending_bold = item_text
                elif pending_bold and not text.startswith("•"):
                    # Linha de detalhe (indentada após bold)
                    if "Responsável:" in text or "Prazo:" in text:
                        resp = ""
                        prazo = ""
                        for part in text.split("  "):
                            part = part.strip()
                            if part.startswith("Responsável:"):
                                resp = part[len("Responsável:"):].strip()
                            elif part.startswith("Prazo:"):
                                prazo = part[len("Prazo:"):].strip()
                        current_items.append({"acao": pending_bold, "responsavel": resp, "prazo": prazo})
                    else:
                        current_items.append({"texto": pending_bold, "detalhe": text})
                    pending_bold = None
                elif pending_bold:
                    current_items.append(pending_bold)
                    pending_bold = None
            elif not current_section and text and result["tipo_reuniao"]:
                # Parágrafo de contexto (vem após os metadados, antes da primeira seção)
                if not result["contexto"]:
                    result["contexto"] = text

    flush_section()
    return result


def download_and_parse(token: str, file_id: str) -> dict:
    import requests, io
    from docx import Document

    headers = {"Authorization": f"Bearer {token}"}
    resp = requests.get(
        f"https://graph.microsoft.com/v1.0/me/drive/items/{file_id}/content",
        headers=headers
    )
    resp.raise_for_status()

    doc = Document(io.BytesIO(resp.content))

    # Detecta formato pelo conteúdo dos parágrafos
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    is_smart = any("Tipo:" in t for t in texts)

    if is_smart:
        return download_and_parse_smart(doc)

    # Legado: documentos antigos no formato A ou B
    result = {
        "titulo": "", "data": "", "time": "", "participantes": [],
        "formato": "legado",
        "acoes_pendentes": [],
        "secoes": [],
    }
    current_section = None
    pending_titulo = None
    legacy_map = {
        "COMBINADOS": "combinados", "PRÓXIMOS PASSOS": "proximos_passos",
        "DECISÕES": "decisoes", "RISCOS": "riscos",
        "PENDÊNCIAS": "pendencias", "REGRAS DE NEGÓCIO": "regras_negocio",
        "O QUE FUNCIONA": "funciona", "O QUE NÃO FUNCIONA": "nao_funciona",
        "ENCAMINHAMENTOS": "encaminhamentos",
    }
    legacy_data = {v: [] for v in legacy_map.values()}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or "═" in text or "─" in text or "RESUMO DE REUNIÃO" in text:
            continue
        if text.startswith("Título:"):
            result["titulo"] = text[len("Título:"):].strip()
        elif text.startswith("Data:"):
            result["data"] = text[len("Data:"):].strip()
        elif text.startswith("Time:"):
            result["time"] = text[len("Time:"):].strip()
        elif text.startswith("Participantes:"):
            result["participantes"] = [p.strip() for p in text[len("Participantes:"):].split(",")]
        elif text.startswith("Resumo gerado"):
            current_section = None
        else:
            for header, key in legacy_map.items():
                if header in text:
                    current_section = key
                    pending_titulo = None
                    break
            else:
                if text.startswith("•") and current_section:
                    item = text.lstrip("•").strip()
                    if item and "Nenhum" not in item:
                        if current_section in ("funciona", "nao_funciona"):
                            pending_titulo = item
                        else:
                            legacy_data[current_section].append(item)
                elif pending_titulo and current_section in ("funciona", "nao_funciona"):
                    legacy_data[current_section].append({"titulo": pending_titulo, "detalhe": text})
                    pending_titulo = None

    # Normaliza para acoes_pendentes (rastreamento semanal)
    for item in legacy_data.get("combinados", []) + legacy_data.get("proximos_passos", []) + legacy_data.get("pendencias", []):
        if isinstance(item, str):
            result["acoes_pendentes"].append({"item": item, "responsavel": "", "prazo": ""})

    # Converte seções legadas para secoes
    for header, key in legacy_map.items():
        items = legacy_data.get(key, [])
        if items:
            result["secoes"].append({"titulo": header.title(), "itens": items})

    return result


def format_html_report(meetings: list, week_start: datetime, week_end: datetime) -> str:
    semana = f"{week_start.strftime('%d/%m')} a {week_end.strftime('%d/%m/%Y')}"

    # Formato A
    combinados_html = ""
    pendencias_html = ""
    proximos_html = ""
    decisoes_html = ""
    # Formato B
    diagnosticos_html = ""

    def meta_label(titulo, team, data):
        return f'<span style="color:#888;font-weight:normal">({team} — {data})</span>'

    def items_list(items):
        if not items:
            return ""
        rows = ""
        for item in items:
            rows += f"<li>{item}</li>"
        return rows

    def diagnostico_block(titulo, team, data, funciona, nao_funciona, encaminhamentos):
        block = f'<div class="diag-block">'
        block += f'<h4>{titulo} {meta_label(titulo, team, data)}</h4>'
        if funciona:
            block += '<p class="diag-label funciona-label">O que funciona</p><ul>'
            for i in funciona:
                t = i.get("titulo", i) if isinstance(i, dict) else i
                d = i.get("detalhe", "") if isinstance(i, dict) else ""
                block += f'<li><strong>{t}</strong>'
                if d:
                    block += f'<br><span class="detalhe">{d}</span>'
                block += '</li>'
            block += '</ul>'
        if nao_funciona:
            block += '<p class="diag-label nao-funciona-label">O que não funciona</p><ul>'
            for i in nao_funciona:
                t = i.get("titulo", i) if isinstance(i, dict) else i
                d = i.get("detalhe", "") if isinstance(i, dict) else ""
                block += f'<li><strong>{t}</strong>'
                if d:
                    block += f'<br><span class="detalhe">{d}</span>'
                block += '</li>'
            block += '</ul>'
        if encaminhamentos:
            block += '<p class="diag-label enc-label">Encaminhamentos</p><ul>'
            for item in encaminhamentos:
                block += f'<li>{item}</li>'
            block += '</ul>'
        block += '</div>'
        return block

    for m in meetings:
        titulo = m.get("titulo") or m.get("name", "")
        team = m.get("team", "")
        data = m.get("data", "")
        fmt = m.get("formato", "padrao")

        if fmt == "funciona_nao_funciona":
            if m.get("funciona") or m.get("nao_funciona"):
                diagnosticos_html += diagnostico_block(
                    titulo, team, data,
                    m.get("funciona", []),
                    m.get("nao_funciona", []),
                    m.get("encaminhamentos", [])
                )
        else:
            if m.get("combinados"):
                combinados_html += f'<h4>{titulo} {meta_label(titulo, team, data)}</h4><ul>{items_list(m["combinados"])}</ul>'
            if m.get("proximos_passos"):
                proximos_html += f'<h4>{titulo} {meta_label(titulo, team, data)}</h4><ul>{items_list(m["proximos_passos"])}</ul>'
            if m.get("pendencias"):
                pendencias_html += f'<h4>{titulo} {meta_label(titulo, team, data)}</h4><ul>{items_list(m["pendencias"])}</ul>'
            if m.get("decisoes"):
                decisoes_html += f'<h4>{titulo} {meta_label(titulo, team, data)}</h4><ul>{items_list(m["decisoes"])}</ul>'

    reunioes_lista = "".join(
        f'<li><strong>{m.get("titulo") or m.get("name","")}</strong> — {m.get("team","")} ({m.get("data","")})'
        f'{"&nbsp;<span class=\'badge-diag\'>diagnóstico</span>" if m.get("formato") == "funciona_nao_funciona" else ""}</li>'
        for m in meetings
    )

    diag_section = ""
    if diagnosticos_html:
        diag_section = f"""
<h2>Diagnósticos — Funciona / Não Funciona</h2>
<div class="section">{diagnosticos_html}</div>
"""

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<style>
  body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 820px; margin: 0 auto; padding: 24px; color: #242424; }}
  h1 {{ color: #1f497d; border-bottom: 2px solid #1f497d; padding-bottom: 8px; }}
  h2 {{ color: #1f497d; margin-top: 32px; border-left: 4px solid #1f497d; padding-left: 12px; }}
  h4 {{ margin: 16px 0 4px; color: #333; }}
  ul {{ margin: 4px 0 12px 0; padding-left: 20px; }}
  li {{ margin-bottom: 6px; line-height: 1.5; }}
  .meta {{ color: #666; font-size: 14px; margin-bottom: 24px; }}
  .section {{ background: #f9f9f9; border-radius: 8px; padding: 16px; margin-bottom: 16px; }}
  .empty {{ color: #999; font-style: italic; }}
  .detalhe {{ color: #555; font-size: 13px; }}
  .diag-block {{ margin-bottom: 24px; border-left: 3px solid #ddd; padding-left: 12px; }}
  .diag-label {{ font-weight: bold; font-size: 12px; text-transform: uppercase; letter-spacing: 0.5px; margin: 10px 0 4px; }}
  .funciona-label {{ color: #117832; }}
  .nao-funciona-label {{ color: #b41e1e; }}
  .enc-label {{ color: #1f497d; }}
  .badge-diag {{ background: #e8f0fe; color: #1f497d; font-size: 11px; padding: 1px 6px; border-radius: 4px; font-weight: normal; }}
</style>
</head>
<body>
<h1>Relatório Semanal de Reuniões</h1>
<p class="meta">Semana de {semana} &nbsp;|&nbsp; Gerado em {datetime.now().strftime("%d/%m/%Y às %H:%M")}</p>

<h2>Reuniões processadas ({len(meetings)})</h2>
<div class="section"><ul>{reunioes_lista or '<li class="empty">Nenhuma reunião processada.</li>'}</ul></div>

{diag_section}

<h2>Combinados</h2>
<div class="section">{combinados_html or '<p class="empty">Nenhum combinado registrado.</p>'}</div>

<h2>Próximos Passos</h2>
<div class="section">{proximos_html or '<p class="empty">Nenhum próximo passo registrado.</p>'}</div>

<h2>Pendências em Aberto</h2>
<div class="section">{pendencias_html or '<p class="empty">Nenhuma pendência registrada.</p>'}</div>

<h2>Decisões Tomadas</h2>
<div class="section">{decisoes_html or '<p class="empty">Nenhuma decisão registrada.</p>'}</div>

</body>
</html>"""
    return html


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", default=str(Path(tempfile.gettempdir()) / "relatorio_semanal.html"))
    parser.add_argument("--json-output", help="Se informado, salva os dados brutos das reuniões como JSON (para análise pelo Claude)")
    args = parser.parse_args()

    ensure_deps()
    token = get_token()

    week_start, week_end = get_week_range()
    print(f"Buscando arquivos da semana: {week_start.strftime('%d/%m')} a {week_end.strftime('%d/%m/%Y')}")

    files = list_week_files(token)
    print(f"Encontrados {len(files)} arquivos")

    meetings = []
    for f in files:
        print(f"  Lendo: {f['name']}")
        try:
            parsed = download_and_parse(token, f["id"])
            parsed["name"] = f["name"]
            parsed["team"] = f["team"]
            meetings.append(parsed)
        except Exception as e:
            print(f"  Aviso: erro ao ler {f['name']}: {e}")

    if args.json_output:
        Path(args.json_output).write_text(
            json.dumps({"semana": f"{week_start.strftime('%d/%m')} a {week_end.strftime('%d/%m/%Y')}", "reunioes": meetings}, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
        print(f"Dados brutos salvos em: {args.json_output}")
        print(json.dumps({"success": True, "meetings": len(meetings), "json_output": args.json_output}))
        return

    html = format_html_report(meetings, week_start, week_end)
    Path(args.output).write_text(html, encoding="utf-8")
    print(f"Relatório salvo em: {args.output}")
    print(json.dumps({"success": True, "meetings": len(meetings), "output": args.output}))


if __name__ == "__main__":
    main()
