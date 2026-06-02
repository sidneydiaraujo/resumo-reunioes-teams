"""
create_docx_smart.py
Cria um .docx a partir do JSON de resumo inteligente (estrutura flexível por teor da reunião).
Uso: python create_docx_smart.py --file resumo.json [--output caminho.docx]
"""

import sys
import json
import subprocess
import argparse
import re
from pathlib import Path
from datetime import datetime


def ensure_deps():
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])


def slugify(text: str) -> str:
    text = re.sub(r'[àáâãäå]', 'a', text.lower())
    text = re.sub(r'[èéêë]', 'e', text)
    text = re.sub(r'[ìíîï]', 'i', text)
    text = re.sub(r'[òóôõö]', 'o', text)
    text = re.sub(r'[ùúûü]', 'u', text)
    text = re.sub(r'[ç]', 'c', text)
    text = re.sub(r'[^a-z0-9\s]', '', text)
    text = re.sub(r'\s+', '_', text.strip())
    return text


def format_date_pt(date_str: str) -> str:
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        meses = ["janeiro","fevereiro","março","abril","maio","junho",
                 "julho","agosto","setembro","outubro","novembro","dezembro"]
        return f"{dt.day} de {meses[dt.month-1]} de {dt.year}"
    except Exception:
        return date_str


TIPO_BADGE = {
    "diagnóstico":   "Diagnóstico",
    "operacional":   "Operacional",
    "estratégico":   "Estratégico",
    "alinhamento":   "Alinhamento",
    "review":        "Review",
    "planning":      "Planning",
    "workshop":      "Workshop",
    "daily":         "Daily",
}


def create_smart_docx(resumo: dict, output_path: str) -> str:
    ensure_deps()
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    BLUE     = RGBColor(31, 73, 125)
    GRAY     = RGBColor(89, 89, 89)
    DIVIDER  = RGBColor(180, 180, 180)
    DETAIL   = RGBColor(85, 85, 85)

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("─" * 60)
        run.font.color.rgb = DIVIDER
        run.font.size = Pt(9)

    def add_meta(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = GRAY
        r2 = p.add_run(value)
        r2.font.size = Pt(10.5)

    def add_section_heading(text):
        add_divider()
        p = doc.add_heading(text, level=2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = BLUE
            run.font.bold = True
            run.font.size = Pt(12)

    def render_item(item):
        """Render one item — string, {texto/detalhe}, or {acao/responsavel/prazo}."""
        if isinstance(item, str):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(3)
            p.add_run(item).font.size = Pt(10.5)

        elif isinstance(item, dict):
            if "acao" in item:
                # Compromisso/ação com dono e prazo
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(item.get("acao", ""))
                r.font.size = Pt(10.5)
                resp = item.get("responsavel", "")
                prazo = item.get("prazo", "")
                if resp or prazo:
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.left_indent = Inches(0.55)
                    p2.paragraph_format.space_after = Pt(5)
                    meta = []
                    if resp:
                        meta.append(f"Responsável: {resp}")
                    if prazo:
                        meta.append(f"Prazo: {prazo}")
                    r2 = p2.add_run("  ".join(meta))
                    r2.font.size = Pt(9.5)
                    r2.font.color.rgb = GRAY

            elif "texto" in item:
                # Item com detalhe explicativo
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(item.get("texto", ""))
                r.bold = True
                r.font.size = Pt(10.5)
                detalhe = item.get("detalhe", "")
                if detalhe:
                    p2 = doc.add_paragraph(detalhe)
                    p2.paragraph_format.left_indent = Inches(0.55)
                    p2.paragraph_format.space_after = Pt(6)
                    p2.runs[0].font.size = Pt(10)
                    p2.runs[0].font.color.rgb = DETAIL

    # ─── Cabeçalho ───
    h = doc.add_heading("RESUMO DE REUNIÃO", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = BLUE

    tipo = resumo.get("tipo_reuniao", "")
    badge = TIPO_BADGE.get(tipo, tipo.capitalize() if tipo else "")

    add_meta("Título", resumo.get("titulo", "Sem título"))
    add_meta("Data", format_date_pt(resumo.get("data", "")))
    add_meta("Time", resumo.get("time", ""))
    if badge:
        add_meta("Tipo", badge)
    add_meta("Participantes", ", ".join(resumo.get("participantes", [])) or "Não informado")

    # ─── Contexto ───
    contexto = resumo.get("contexto", "")
    if contexto:
        add_divider()
        p = doc.add_paragraph(contexto)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.italic = True
            run.font.color.rgb = GRAY

    # ─── Seções inteligentes ───
    for secao in resumo.get("secoes", []):
        titulo_secao = secao.get("titulo", "")
        itens = secao.get("itens", [])
        if not titulo_secao and not itens:
            continue

        add_section_heading(titulo_secao.upper())
        if not itens:
            p = doc.add_paragraph("Nenhum item registrado.")
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
            continue
        for item in itens:
            render_item(item)

    # ─── Ações Pendentes ───
    acoes = resumo.get("acoes_pendentes", [])
    if acoes:
        add_section_heading("AÇÕES PENDENTES")
        for item in acoes:
            if isinstance(item, dict) and "item" in item:
                render_item({
                    "acao": item.get("item", ""),
                    "responsavel": item.get("responsavel", ""),
                    "prazo": item.get("prazo", "")
                })
            else:
                render_item(item)

    # ─── Contextos Anteriores ───
    contextos = resumo.get("contextos_anteriores", {})
    resolvidos = contextos.get("resolvidos", [])
    persistentes = contextos.get("persistentes", [])

    if resolvidos or persistentes:
        add_section_heading("CONTEXTOS DE REUNIÕES ANTERIORES")

        GREEN  = RGBColor(0, 112, 0)
        ORANGE = RGBColor(197, 90, 17)

        if resolvidos:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run("✅  Resolvidos nesta reunião")
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = GREEN
            for ctx in resolvidos:
                p2 = doc.add_paragraph(style='List Bullet')
                p2.paragraph_format.left_indent = Inches(0.3)
                p2.paragraph_format.space_after = Pt(2)
                texto = ctx.get("item", "")
                como  = ctx.get("como", "")
                data_r = ctx.get("resolvido_em", "")
                p2.add_run(texto).font.size = Pt(10.5)
                if como or data_r:
                    p3 = doc.add_paragraph()
                    p3.paragraph_format.left_indent = Inches(0.55)
                    p3.paragraph_format.space_after = Pt(5)
                    meta = []
                    if data_r:
                        meta.append(f"Resolvido em: {data_r}")
                    if como:
                        meta.append(como)
                    r3 = p3.add_run("  ".join(meta))
                    r3.font.size = Pt(9.5)
                    r3.font.color.rgb = GREEN

        if persistentes:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run("⏳  Ainda em aberto")
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = ORANGE
            for ctx in persistentes:
                p2 = doc.add_paragraph(style='List Bullet')
                p2.paragraph_format.left_indent = Inches(0.3)
                p2.paragraph_format.space_after = Pt(2)
                texto = ctx.get("item", "")
                status = ctx.get("status_atual", "")
                desde  = ctx.get("aberto_desde", "")
                p2.add_run(texto).font.size = Pt(10.5)
                if status or desde:
                    p3 = doc.add_paragraph()
                    p3.paragraph_format.left_indent = Inches(0.55)
                    p3.paragraph_format.space_after = Pt(5)
                    meta = []
                    if desde:
                        meta.append(f"Em aberto desde: {desde}")
                    if status:
                        meta.append(f"Status: {status}")
                    r3 = p3.add_run("  ".join(meta))
                    r3.font.size = Pt(9.5)
                    r3.font.color.rgb = ORANGE

    # ─── IDs Jira mencionados ───
    jira_ids = resumo.get("jira_ids_mencionados", [])
    if jira_ids:
        add_divider()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("IDs Jira mencionados:  ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = GRAY
        r2 = p.add_run("  ·  ".join(jira_ids))
        r2.font.size = Pt(10)
        r2.font.color.rgb = BLUE

    # ─── Rodapé ───
    add_divider()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Resumo gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.italic = True

    doc.save(output_path)
    return output_path


def build_filename(resumo: dict) -> str:
    data = resumo.get("data", datetime.now().strftime("%Y-%m-%d"))
    titulo = slugify(resumo.get("titulo", "reuniao"))[:40]
    time_ = slugify(resumo.get("time", "geral"))[:20]
    return f"{data}_{titulo}_{time_}.docx"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("json_input", nargs="?", help="JSON inline do resumo")
    parser.add_argument("--file", help="Caminho para arquivo .json")
    parser.add_argument("--output", help="Caminho de saída do .docx")
    args = parser.parse_args()

    if args.file:
        with open(args.file, encoding="utf-8") as f:
            resumo = json.load(f)
    elif args.json_input:
        resumo = json.loads(args.json_input)
    else:
        resumo = json.loads(sys.stdin.read())

    filename = build_filename(resumo)
    output = args.output or str(Path.home() / "Downloads" / filename)

    result_path = create_smart_docx(resumo, output)
    print(json.dumps({"success": True, "path": result_path, "filename": filename}, ensure_ascii=False))


if __name__ == "__main__":
    main()
