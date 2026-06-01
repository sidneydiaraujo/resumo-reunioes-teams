"""
create_docx_diagnostico.py
Cria um .docx no formato Funciona / Não Funciona a partir do JSON de resumo diagnóstico.
Uso: python create_docx_diagnostico.py --file resumo.json [--output caminho.docx]
"""

import sys
import json
import subprocess
import argparse
import re
from pathlib import Path
from datetime import datetime


def ensure_deps():
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])


def slugify(text: str) -> str:
    text = re.sub(r'[àáâãäå]', 'a', text.lower())
    text = re.sub(r'[èéêë]', 'e', text)
    text = re.sub(r'[ìíîï]', 'i', text)
    text = re.sub(r'[òóôõö]', 'o', text)
    text = re.sub(r'[ùúûü]', 'u', text)
    text = re.sub(r'[ç]', 'c', text)
    text = re.sub(r'[^a-z0-9\s]', '', text)
    text = re.sub(r'\s+', '_', text.strip())
    return text


def format_date_pt(date_str: str) -> str:
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        meses = ["janeiro","fevereiro","março","abril","maio","junho",
                 "julho","agosto","setembro","outubro","novembro","dezembro"]
        return f"{dt.day} de {meses[dt.month-1]} de {dt.year}"
    except Exception:
        return date_str


def create_diagnostico_docx(resumo: dict, output_path: str) -> str:
    ensure_deps()
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("─" * 60)
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.size = Pt(9)

    def add_meta(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(10.5)
        run_label.font.color.rgb = RGBColor(89, 89, 89)
        run_value = p.add_run(value)
        run_value.font.size = Pt(10.5)

    def add_section_heading(text, color):
        add_divider()
        p = doc.add_heading(text, level=2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
            run.font.bold = True
            run.font.size = Pt(12)
        doc.add_paragraph()

    def add_item(titulo, detalhe):
        # Título do item em bold
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {titulo}")
        r.bold = True
        r.font.size = Pt(11)
        # Detalhe em cinza abaixo
        p2 = doc.add_paragraph(detalhe)
        p2.paragraph_format.left_indent = Inches(0.35)
        p2.paragraph_format.space_after = Pt(8)
        p2.runs[0].font.size = Pt(10)
        p2.runs[0].font.color.rgb = RGBColor(68, 68, 68)

    def add_encaminhamento(item, responsavel, prazo):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"• {item}")
        r.bold = True
        r.font.size = Pt(11)
        p2 = doc.add_paragraph(f"Responsável: {responsavel}   |   Prazo: {prazo}")
        p2.paragraph_format.left_indent = Inches(0.35)
        p2.paragraph_format.space_after = Pt(8)
        p2.runs[0].font.size = Pt(10)
        p2.runs[0].font.color.rgb = RGBColor(68, 68, 68)

    # ─── Cabeçalho ───
    titulo_h = doc.add_heading("RESUMO DE REUNIÃO", level=1)
    titulo_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo_h.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 73, 125)

    add_meta("Título", resumo.get("titulo", "Sem título"))
    add_meta("Data", format_date_pt(resumo.get("data", "")))
    add_meta("Time", resumo.get("time", ""))
    add_meta("Participantes", ", ".join(resumo.get("participantes", [])) or "Não informado")

    # ─── O que funciona ───
    funciona = resumo.get("funciona", [])
    add_section_heading("O QUE FUNCIONA", color=(17, 120, 54))
    if funciona:
        for item in funciona:
            add_item(item.get("titulo", ""), item.get("detalhe", ""))
    else:
        p = doc.add_paragraph("Nenhum ponto registrado.")
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # ─── O que não funciona ───
    nao_funciona = resumo.get("nao_funciona", [])
    add_section_heading("O QUE NÃO FUNCIONA", color=(180, 30, 30))
    if nao_funciona:
        for item in nao_funciona:
            add_item(item.get("titulo", ""), item.get("detalhe", ""))
    else:
        p = doc.add_paragraph("Nenhum ponto registrado.")
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # ─── Encaminhamentos (somente se houver) ───
    encaminhamentos = resumo.get("encaminhamentos", [])
    if encaminhamentos:
        add_section_heading("ENCAMINHAMENTOS", color=(31, 73, 125))
        for enc in encaminhamentos:
            add_encaminhamento(
                enc.get("item", ""),
                enc.get("responsavel", "A definir"),
                enc.get("prazo", "A definir")
            )

    # ─── Rodapé ───
    add_divider()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Resumo gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.italic = True

    doc.save(output_path)
    return output_path


def build_filename(resumo: dict) -> str:
    data = resumo.get("data", datetime.now().strftime("%Y-%m-%d"))
    titulo = slugify(resumo.get("titulo", "Reuniao"))[:40]
    time_ = slugify(resumo.get("time", "Geral"))[:20]
    return f"{data}_{titulo}_{time_}.docx"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("json_input", nargs="?", help="JSON inline do resumo")
    parser.add_argument("--file", help="Caminho para arquivo .json com o resumo")
    parser.add_argument("--output", help="Caminho de saída do .docx")
    args = parser.parse_args()

    if args.file:
        with open(args.file, encoding="utf-8") as f:
            resumo = json.load(f)
    elif args.json_input:
        resumo = json.loads(args.json_input)
    else:
        raw = sys.stdin.read()
        resumo = json.loads(raw)

    filename = build_filename(resumo)
    output = args.output or str(Path.home() / "Downloads" / filename)

    result_path = create_diagnostico_docx(resumo, output)
    print(json.dumps({"success": True, "path": result_path, "filename": filename}, ensure_ascii=False))


if __name__ == "__main__":
    main()
