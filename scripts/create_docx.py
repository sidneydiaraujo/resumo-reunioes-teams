"""
create_docx.py
Cria um arquivo .docx formatado a partir do JSON de resumo de reunião.
Uso: python create_docx.py '<json_resumo>' [--output caminho.docx]
     python create_docx.py --file resumo.json [--output caminho.docx]
"""

import sys
import json
import subprocess
import argparse
import re
from pathlib import Path
from datetime import datetime


def ensure_deps():
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])


def slugify(text: str) -> str:
    text = re.sub(r'[àáâãäå]', 'a', text.lower())
    text = re.sub(r'[èéêë]', 'e', text)
    text = re.sub(r'[ìíîï]', 'i', text)
    text = re.sub(r'[òóôõö]', 'o', text)
    text = re.sub(r'[ùúûü]', 'u', text)
    text = re.sub(r'[ç]', 'c', text)
    text = re.sub(r'[^a-z0-9\s]', '', text)
    text = re.sub(r'\s+', '_', text.strip())
    return text


def format_date_pt(date_str: str) -> str:
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        meses = ["janeiro","fevereiro","março","abril","maio","junho",
                 "julho","agosto","setembro","outubro","novembro","dezembro"]
        return f"{dt.day} de {meses[dt.month-1]} de {dt.year}"
    except Exception:
        return date_str


def create_summary_docx(resumo: dict, output_path: str) -> str:
    ensure_deps()
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # margens
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    def add_heading(text, level=1, color=(31, 73, 125)):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
            run.font.bold = True
            if level == 1:
                run.font.size = Pt(16)
            elif level == 2:
                run.font.size = Pt(12)
        return p

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("─" * 60)
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.size = Pt(9)

    def add_meta(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(10.5)
        run_label.font.color.rgb = RGBColor(89, 89, 89)
        run_value = p.add_run(value)
        run_value.font.size = Pt(10.5)

    def add_section(titulo, itens, formatter):
        add_divider()
        add_heading(titulo, level=2, color=(31, 73, 125))
        if not itens:
            p = doc.add_paragraph()
            run = p.add_run(get_empty_msg(titulo))
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.italic = True
            run.font.size = Pt(10)
            return
        for item in itens:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(3)
            text = formatter(item)
            for part in text.split("**"):
                pass
            # escreve o texto formatado
            segments = re.split(r'(\*\*[^*]+\*\*)', text)
            for seg in segments:
                bold_match = re.match(r'\*\*([^*]+)\*\*', seg)
                if bold_match:
                    run = p.add_run(bold_match.group(1))
                    run.bold = True
                    run.font.size = Pt(10.5)
                elif seg:
                    run = p.add_run(seg)
                    run.font.size = Pt(10.5)

    def get_empty_msg(titulo):
        msgs = {
            "REGRAS DE NEGÓCIO": "Nenhuma regra de negócio mencionada nesta reunião.",
            "COMBINADOS DA REUNIÃO": "Nenhum combinado registrado.",
            "PRÓXIMOS PASSOS": "Nenhum próximo passo registrado.",
            "DECISÕES TOMADAS": "Nenhuma decisão formal registrada.",
            "RISCOS E IMPEDIMENTOS": "Nenhum risco ou impedimento levantado.",
            "PENDÊNCIAS": "Nenhuma pendência registrada.",
        }
        return msgs.get(titulo, "Nenhum item registrado.")

    # ─── Cabeçalho ───
    titulo_h = doc.add_heading("RESUMO DE REUNIÃO", level=1)
    titulo_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo_h.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 73, 125)

    add_meta("Título", resumo.get("titulo", "Sem título"))
    add_meta("Data", format_date_pt(resumo.get("data", "")))
    add_meta("Time", resumo.get("time", ""))
    add_meta("Participantes", ", ".join(resumo.get("participantes", [])) or "Não informado")

    # ─── Seções ───
    add_section(
        "REGRAS DE NEGÓCIO",
        resumo.get("regras_negocio", []),
        lambda x: x if isinstance(x, str) else str(x)
    )

    add_section(
        "COMBINADOS DA REUNIÃO",
        resumo.get("combinados", []),
        lambda x: f"{x.get('item','')} — **Responsável:** {x.get('responsavel','A definir')} | **Prazo:** {x.get('prazo','A definir')}"
        if isinstance(x, dict) else str(x)
    )

    add_section(
        "PRÓXIMOS PASSOS",
        resumo.get("proximos_passos", []),
        lambda x: f"{x.get('acao','')} — **Responsável:** {x.get('responsavel','A definir')} | **Prazo:** {x.get('prazo','A definir')}"
        if isinstance(x, dict) else str(x)
    )

    add_section(
        "DECISÕES TOMADAS",
        resumo.get("decisoes", []),
        lambda x: x if isinstance(x, str) else str(x)
    )

    add_section(
        "RISCOS E IMPEDIMENTOS",
        resumo.get("riscos", []),
        lambda x: f"{x.get('item','')} — **Impacto:** {x.get('impacto','Não classificado')}"
        if isinstance(x, dict) else str(x)
    )

    add_section(
        "PENDÊNCIAS",
        resumo.get("pendencias", []),
        lambda x: f"{x.get('item','')} — **Responsável:** {x.get('responsavel','A definir')} | **Prazo de resolução:** {x.get('prazo','A definir')}"
        if isinstance(x, dict) else str(x)
    )

    # ─── Rodapé ───
    add_divider()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Resumo gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.italic = True

    doc.save(output_path)
    return output_path


def build_filename(resumo: dict) -> str:
    data = resumo.get("data", datetime.now().strftime("%Y-%m-%d"))
    titulo = slugify(resumo.get("titulo", "Reuniao"))[:40]
    time_ = slugify(resumo.get("time", "Geral"))[:20]
    return f"{data}_{titulo}_{time_}.docx"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("json_input", nargs="?", help="JSON inline do resumo")
    parser.add_argument("--file", help="Caminho para arquivo .json com o resumo")
    parser.add_argument("--output", help="Caminho de saída do .docx")
    args = parser.parse_args()

    if args.file:
        with open(args.file, encoding="utf-8") as f:
            resumo = json.load(f)
    elif args.json_input:
        resumo = json.loads(args.json_input)
    else:
        raw = sys.stdin.read()
        resumo = json.loads(raw)

    filename = build_filename(resumo)
    output = args.output or str(Path.home() / "Downloads" / filename)

    result_path = create_summary_docx(resumo, output)
    print(json.dumps({"success": True, "path": result_path, "filename": filename}, ensure_ascii=False))


if __name__ == "__main__":
    main()
