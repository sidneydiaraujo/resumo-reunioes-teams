"""
parse_transcript.py
Extrai texto estruturado de transcrições do Teams (.vtt ou .docx).
Uso: python parse_transcript.py <caminho_do_arquivo>
Retorna JSON com lista de falas: [{"speaker": "...", "text": "...", "time": "..."}]
"""

import sys
import json
import re
import subprocess
from pathlib import Path


def ensure_deps():
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])


def parse_vtt(content: str) -> list[dict]:
    falas = []
    current_speaker = None
    current_text_parts = []
    current_time = None

    lines = content.strip().splitlines()
    i = 0
    # pula cabeçalho WEBVTT
    while i < len(lines) and not re.match(r'\d{2}:\d{2}', lines[i]):
        i += 1

    while i < len(lines):
        line = lines[i].strip()

        # linha de timestamp: 00:01:23.000 --> 00:01:28.000
        ts_match = re.match(r'(\d{2}:\d{2}:\d{2}[\.,]\d{3})\s*-->', line)
        if ts_match:
            # salva fala anterior
            if current_speaker and current_text_parts:
                falas.append({
                    "speaker": current_speaker.strip(),
                    "text": " ".join(current_text_parts).strip(),
                    "time": current_time or ""
                })
                current_text_parts = []

            current_time = ts_match.group(1)
            i += 1
            continue

        # linha vazia: separador de bloco
        if not line:
            i += 1
            continue

        # número de sequência (só dígitos)
        if re.match(r'^\d+$', line):
            i += 1
            continue

        # tag <v NomeFalante>texto</v>
        v_match = re.match(r'<v\s+([^>]+)>(.*?)(?:</v>)?$', line, re.IGNORECASE)
        if v_match:
            if current_speaker and current_speaker != v_match.group(1).strip() and current_text_parts:
                falas.append({
                    "speaker": current_speaker.strip(),
                    "text": " ".join(current_text_parts).strip(),
                    "time": current_time or ""
                })
                current_text_parts = []
            current_speaker = v_match.group(1).strip()
            text = re.sub(r'<[^>]+>', '', v_match.group(2)).strip()
            if text:
                current_text_parts.append(text)
            i += 1
            continue

        # padrão "Nome Falante: texto"
        colon_match = re.match(r'^([A-ZÀ-Ú][^:]{2,40}):\s+(.+)$', line)
        if colon_match:
            if current_speaker and current_speaker != colon_match.group(1).strip() and current_text_parts:
                falas.append({
                    "speaker": current_speaker.strip(),
                    "text": " ".join(current_text_parts).strip(),
                    "time": current_time or ""
                })
                current_text_parts = []
            current_speaker = colon_match.group(1).strip()
            text = re.sub(r'<[^>]+>', '', colon_match.group(2)).strip()
            if text:
                current_text_parts.append(text)
            i += 1
            continue

        # continuação de fala (remove tags HTML)
        cleaned = re.sub(r'<[^>]+>', '', line).strip()
        if cleaned and current_speaker:
            current_text_parts.append(cleaned)

        i += 1

    # última fala
    if current_speaker and current_text_parts:
        falas.append({
            "speaker": current_speaker.strip(),
            "text": " ".join(current_text_parts).strip(),
            "time": current_time or ""
        })

    return falas


def parse_docx(filepath: str) -> list[dict]:
    ensure_deps()
    import docx

    doc = docx.Document(filepath)
    falas = []

    # Tenta extrair de tabela (formato padrão do Teams)
    for table in doc.tables:
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        has_speaker = any(h in headers for h in ['speaker', 'falante', 'participante', 'nome'])
        has_text = any(h in headers for h in ['text', 'texto', 'transcrição', 'transcription', 'content', 'message'])

        if has_speaker and has_text:
            speaker_idx = next(
                (i for i, h in enumerate(headers) if h in ['speaker', 'falante', 'participante', 'nome']), 0
            )
            text_idx = next(
                (i for i, h in enumerate(headers) if h in ['text', 'texto', 'transcrição', 'transcription', 'content', 'message']), 1
            )
            time_idx = next(
                (i for i, h in enumerate(headers) if 'time' in h or 'hora' in h or 'tempo' in h), None
            )

            for row in table.rows[1:]:
                cells = row.cells
                if len(cells) <= max(speaker_idx, text_idx):
                    continue
                speaker = cells[speaker_idx].text.strip()
                text = cells[text_idx].text.strip()
                time_val = cells[time_idx].text.strip() if time_idx is not None and time_idx < len(cells) else ""
                if speaker and text:
                    falas.append({"speaker": speaker, "text": text, "time": time_val})
            if falas:
                return falas

    # Fallback: extrai parágrafos como texto livre
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # tenta detectar padrão "Nome: fala"
        colon_match = re.match(r'^([A-ZÀ-Ú][^:]{2,40}):\s+(.+)$', text)
        if colon_match:
            falas.append({
                "speaker": colon_match.group(1).strip(),
                "text": colon_match.group(2).strip(),
                "time": ""
            })
        else:
            falas.append({"speaker": "Desconhecido", "text": text, "time": ""})

    return falas


def consolidate(falas: list[dict]) -> str:
    """Agrupa falas consecutivas do mesmo falante para reduzir ruído."""
    if not falas:
        return ""

    consolidated = []
    current = falas[0].copy()

    for fala in falas[1:]:
        if fala["speaker"] == current["speaker"]:
            current["text"] += " " + fala["text"]
        else:
            consolidated.append(current)
            current = fala.copy()
    consolidated.append(current)

    lines = []
    for f in consolidated:
        lines.append(f"[{f['time']}] {f['speaker']}: {f['text']}" if f['time'] else f"{f['speaker']}: {f['text']}")

    return "\n\n".join(lines)


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Uso: python parse_transcript.py <caminho_arquivo>"}))
        sys.exit(1)

    filepath = sys.argv[1]
    path = Path(filepath)

    if not path.exists():
        print(json.dumps({"error": f"Arquivo não encontrado: {filepath}"}))
        sys.exit(1)

    ext = path.suffix.lower()

    if ext == ".vtt":
        content = path.read_text(encoding="utf-8", errors="replace")
        falas = parse_vtt(content)
    elif ext in (".docx", ".doc"):
        falas = parse_docx(str(path))
    else:
        print(json.dumps({"error": f"Formato não suportado: {ext}. Use .vtt ou .docx"}))
        sys.exit(1)

    speakers = list(dict.fromkeys(f["speaker"] for f in falas if f["speaker"] != "Desconhecido"))
    transcript_text = consolidate(falas)

    result = {
        "speakers": speakers,
        "total_falas": len(falas),
        "transcript": transcript_text
    }

    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
